import json
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_FOLDER = "json"                   # <- your flat input folder
OUTPUT_FOLDER = "output_results/doc"    # <- output folder

def json_to_docx(json_data, output_file):
    doc = Document()
    json_str = json.dumps(json_data, indent=4, ensure_ascii=False)

    p = doc.add_paragraph()
    run = p.add_run(json_str)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # Force Calibri in Word (XML font mapping)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:cs'), 'Calibri')
    rPr.append(rFonts)

    doc.save(output_file)
    print(f"✅ Saved: {output_file}")

def main():
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    for filename in os.listdir(INPUT_FOLDER):
        if not filename.lower().endswith(".json"):
            continue

        input_path = os.path.join(INPUT_FOLDER, filename)
        output_name = os.path.splitext(filename)[0] + ".docx"
        output_path = os.path.join(OUTPUT_FOLDER, output_name)

        try:
            with open(input_path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            print(f"⚠️ Skipping invalid JSON: {filename} ({e})")
            continue

        json_to_docx(data, output_path)

if __name__ == "__main__":
    main()
